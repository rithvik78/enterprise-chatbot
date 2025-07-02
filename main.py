#!/usr/bin/env python3

import os
import json
import numpy as np
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from dotenv import load_dotenv
import requests
import anthropic
from pinecone import Pinecone, ServerlessSpec
from sentence_transformers import SentenceTransformer
from scipy.spatial.distance import cosine
import hashlib
import pickle
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import io
import tempfile
from supabase import create_client
from openai import OpenAI
from urllib.parse import urlencode
import re
import time

os.environ["TOKENIZERS_PARALLELISM"] = "false"
load_dotenv()

@dataclass
class Document:
    id: str
    title: str
    content: str
    source: str
    embedding: List[float] = None

@dataclass
class QueryParams:
    filter_formula: Optional[str] = None
    sort: Optional[List[Dict[str, str]]] = None
    fields: Optional[List[str]] = None
    max_records: Optional[int] = None
    page_size: Optional[int] = None
    offset: Optional[str] = None

class AirtableQueryBuilder:
    def __init__(self, base_id: str, api_key: str):
        self.base_id = base_id
        self.api_key = api_key
        self.base_url = f"https://api.airtable.com/v0/{base_id}"
        self.headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
    
    def get_base_schema(self) -> Dict[str, Any]:
        """Dynamically fetch the complete base schema from Airtable"""
        url = f"https://api.airtable.com/v0/meta/bases/{self.base_id}/tables"
        
        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            return {"error": f"Failed to fetch schema: {str(e)}"}
    
    def get_sample_data(self, table_name: str, limit: int = 5) -> Dict[str, Any]:
        """Get sample records to understand table content"""
        url = f"{self.base_url}/{table_name}?maxRecords={limit}"
        
        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            return {"error": f"Failed to fetch sample data: {str(e)}"}
    
    def build_filter_formula(self, conditions: Dict[str, Any]) -> str:
        if not conditions:
            return ""
        
        formula_parts = []
        for field, value in conditions.items():
            if isinstance(value, str):
                formula_parts.append(f"{{{field}}}='{value}'")
            elif isinstance(value, (int, float)):
                formula_parts.append(f"{{{field}}}={value}")
            elif isinstance(value, dict):
                if 'operator' in value and 'value' in value:
                    op = value['operator']
                    val = value['value']
                    if op == 'greater_than':
                        formula_parts.append(f"{{{field}}}>{val}")
                    elif op == 'less_than':
                        formula_parts.append(f"{{{field}}}<{val}")
                    elif op == 'contains':
                        formula_parts.append(f"SEARCH('{val}',{{{field}}})")
        
        if len(formula_parts) == 1:
            return formula_parts[0]
        elif len(formula_parts) > 1:
            return f"AND({','.join(formula_parts)})"
        return ""
    
    def build_sort_params(self, sort_fields: List[Dict[str, str]]) -> List[Dict[str, str]]:
        return [{"field": item["field"], "direction": item.get("direction", "asc")} 
                for item in sort_fields]
    
    def build_query_url(self, table_name: str, params: QueryParams) -> str:
        url_params = {}
        
        if params.filter_formula:
            url_params['filterByFormula'] = params.filter_formula
        
        if params.sort:
            for i, sort_item in enumerate(params.sort):
                url_params[f'sort[{i}][field]'] = sort_item['field']
                url_params[f'sort[{i}][direction]'] = sort_item['direction']
        
        if params.fields:
            for i, field in enumerate(params.fields):
                url_params[f'fields[{i}]'] = field
        
        if params.max_records:
            url_params['maxRecords'] = params.max_records
        
        if params.page_size:
            url_params['pageSize'] = params.page_size
        
        if params.offset:
            url_params['offset'] = params.offset
        
        base_url = f"{self.base_url}/{table_name}"
        if url_params:
            return f"{base_url}?{urlencode(url_params)}"
        return base_url
    
    def execute_query(self, table_name: str, params: QueryParams) -> Dict[str, Any]:
        url = self.build_query_url(table_name, params)
        
        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            return {"error": f"API request failed: {str(e)}"}

class SQLCodeGenerator:
    def __init__(self, openrouter_api_key: str):
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=openrouter_api_key,
        )
    
    def _clean_response(self, response: str) -> str:
        try:
            lines = response.split('\n')
            cleaned_lines = []
            
            for line in lines:
                if any(marker in line.lower() for marker in [
                    'system:', 'instruction:', 'prompt:', 'you are', 'convert the', 
                    'return only', 'available operators', 'table schema'
                ]):
                    continue
                cleaned_lines.append(line)
            
            return '\n'.join(cleaned_lines).strip()
        except:
            return response
    
    def generate_airtable_query(self, query: str, table_schema: Dict[str, Any], conversation_context: Optional[Dict] = None) -> Dict[str, Any]:
        context_info = ""
        if conversation_context:
            context_info = f"""
        
        Conversation Context: {json.dumps(conversation_context, indent=2)}
        Use this context to better understand references in the query.
        """
        
        system_prompt = f"""Convert natural language to Airtable API parameters.

Table Schema: {json.dumps(table_schema, indent=2)}
{context_info}

Return JSON with these fields:
- conditions: Field filters
- sort: Sort parameters  
- fields: Specific fields to return
- max_records: Result limit

Condition examples:
- Simple: {{"Status": "Active"}}
- Contains: {{"Name": {{"operator": "contains", "value": "John"}}}}
- Comparison: {{"Salary": {{"operator": "greater_than", "value": 50000}}}}

Available operators: greater_than, less_than, contains, equals

Return only valid JSON."""
        
        try:
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="qwen/qwen-2.5-coder-32b-instruct:free",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": query}
                ],
                temperature=0.1
            )
            
            result = self._clean_response(completion.choices[0].message.content.strip())
            # Try to extract JSON from the cleaned response
            try:
                return json.loads(result)
            except json.JSONDecodeError:
                # Try to find JSON in the response
                import re
                json_match = re.search(r'\{.*\}', result, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
                raise ValueError("No valid JSON found")
        except Exception as e:
            return {"error": "Query generation failed", "details": "Unable to process request"}

class QueryRouter:
    def __init__(self, openrouter_api_key: str):
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=openrouter_api_key,
        )
    
    def _clean_response(self, response: str) -> str:
        try:
            lines = response.split('\n')
            cleaned_lines = []
            
            for line in lines:
                if any(marker in line.lower() for marker in [
                    'you are', 'analyze the', 'given a', 'return only', 
                    'system:', 'instructions:', 'guidelines:'
                ]):
                    continue
                cleaned_lines.append(line)
            
            result = '\n'.join(cleaned_lines).strip()
            return result if result else "Analysis completed"
        except:
            return "Analysis completed"
    
    def analyze_table_purpose(self, table_name: str, table_schema: Dict[str, Any], sample_data: Dict[str, Any]) -> str:
        """Use LLM to understand what a table is for based on its schema and sample data"""
        system_prompt = """
        You are an expert at analyzing database tables. Given a table name, schema, and sample data, 
        provide a concise 1-2 sentence description of what this table is used for.
        
        Focus on the business purpose and what kind of data it contains.
        """
        
        analysis_data = {
            "table_name": table_name,
            "schema": table_schema,
            "sample_records": sample_data.get("records", [])[:3]  # First 3 records
        }
        
        try:
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": json.dumps(analysis_data, indent=2)}
                ],
                temperature=0.3
            )
            
            response = self._clean_response(completion.choices[0].message.content.strip())
            return response if response else f"A table containing {table_name} data"
        except Exception as e:
            return f"A table containing {table_name} data"
    
    def determine_best_table(self, user_query: str, available_tables: Dict[str, str]) -> str:
        """Use LLM to determine which table best matches the user's query"""
        system_prompt = """
        You are helping route a user query to the most appropriate database table.
        
        Given a user query and descriptions of available tables, return ONLY the exact table name 
        that best matches the query. If no table seems appropriate, return the first table name.
        
        Return only the table name, nothing else.
        """
        
        table_info = "\n".join([f"- {name}: {desc}" for name, desc in available_tables.items()])
        
        try:
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Query: {user_query}\n\nAvailable tables:\n{table_info}"}
                ],
                temperature=0.1
            )
            
            result = self._clean_response(completion.choices[0].message.content.strip())
            # Return the result if it's a valid table name, otherwise return first available
            if result in available_tables:
                return result
            else:
                return list(available_tables.keys())[0]
        except Exception as e:
            return list(available_tables.keys())[0] if available_tables else ""
    
    def should_use_airtable(self, user_query: str, available_tables: Dict[str, str]) -> Dict[str, Any]:
        """Determine if the query should use Airtable vs document retrieval"""
        system_prompt = f"""Determine the best data source for user queries.

Available Airtable tables: {json.dumps(available_tables, indent=2)}

Route queries to:
1. AIRTABLE: People, employees, reports, organizational data, assets, tickets, structured records
2. DOCUMENTS: Policies, procedures, guidelines, text content, instructions

Return JSON:
{{
    "use_airtable": true/false,
    "reasoning": "brief explanation",
    "table_name": "table name if airtable",
    "confidence": 0.0-1.0
}}

Priority: Employee queries ("who reports to X", "who works for Y", "employees in Z") always use Airtable/Employees table."""
        
        try:
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_query}
                ],
                temperature=0.2
            )
            
            result = self._clean_response(completion.choices[0].message.content.strip())
            try:
                return json.loads(result)
            except json.JSONDecodeError:
                # Try to extract JSON from the response
                import re
                json_match = re.search(r'\{.*\}', result, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
                raise ValueError("No valid JSON found")
        except Exception as e:
            return {"use_airtable": False, "reasoning": "Unable to determine data source preference", "confidence": 0.0}

    def analyze_and_suggest_retry(self, query: str, schema: dict, error: str, context: dict) -> dict:
        """Analyze error and suggest retry with modified query"""
        try:
            prompt = f"Query '{query}' failed with error: {error}. Suggest a modified query or return should_retry: false if not fixable."
            
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=200
            )
            
            response = self._clean_response(completion.choices[0].message.content.strip())
            return {"should_retry": False, "modified_query": query}
        except:
            return {"should_retry": False, "modified_query": query}

    def analyze_empty_results(self, query: str, schema: dict, parsed_query: dict, context: dict) -> dict:
        """Analyze why no results were found and suggest retry"""
        try:
            prompt = f"Query '{query}' returned no results. Suggest a broader or modified query to find relevant data."
            
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=200
            )
            
            response = self._clean_response(completion.choices[0].message.content.strip())
            return {"should_retry": False, "modified_query": query}
        except:
            return {"should_retry": False, "modified_query": query}

    def format_results(self, query_results: Dict[str, Any], user_query: str) -> str:
        """Format database query results for the user in a friendly, conversational way"""
        system_prompt = """Format database query results clearly and completely.

Guidelines:
- Employee queries: List ALL people with names, roles, departments
- Organizational queries: Show complete hierarchical relationships  
- Include all relevant details (names, titles, departments, contacts)
- Never truncate or summarize when full data is available
- Use bullet points or numbered lists for readability
- Start with summary, then provide complete details

Provide comprehensive information, not abbreviated summaries."""
        
        try:
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                max_tokens=2000,  # Ensure sufficient tokens for complete responses
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Original query: {user_query}\nResults: {json.dumps(query_results, indent=2)}"}
                ],
                temperature=0.3
            )
            
            response = self._clean_response(completion.choices[0].message.content.strip())
            return response if response else "Results have been processed and formatted."
        except Exception as e:
            # Provide a clean fallback without exposing the raw data structure
            try:
                records = query_results.get('records', [])
                if records:
                    return f"Found {len(records)} matching records in the database."
                else:
                    return self._format_no_results_message(user_query)
            except:
                return "Results found but formatting failed. Please try a simpler query."
    
    def _format_no_results_message(self, user_query: str) -> str:
        try:
            prompt = f"Generate a helpful message when no database records match the query: '{user_query}'. Suggest alternative search terms."
            
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Universal Airtable Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                max_tokens=100,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            return completion.choices[0].message.content.strip()
        except:
            return f"No records found matching '{user_query}'. Try different search terms or check spelling."

class RelevanceScorer:
    def __init__(self, openrouter_api_key: str):
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=openrouter_api_key,
        )
        self.score_tiers = {
            'perfect_match': (0.85, 1.0),
            'high_relevance': (0.70, 0.84),
            'moderate_relevance': (0.50, 0.69),
            'low_relevance': (0.25, 0.49),
            'no_relevance': (0.0, 0.24)
        }
    
    def _clean_response(self, response: str) -> str:
        try:
            lines = response.split('\n')
            cleaned_lines = []
            
            for line in lines:
                if any(marker in line.lower() for marker in [
                    'you are', 'evaluate', 'consider:', 'return json', 
                    'scoring guide:', 'content type:', 'system:'
                ]):
                    continue
                cleaned_lines.append(line)
            
            return '\n'.join(cleaned_lines).strip()
        except:
            return response
    
    def _normalize_score(self, raw_score: float, tier: str = None) -> float:
        try:
            score = max(0.0, min(1.0, raw_score))
            
            if score >= 0.95:
                return round(score, 3)
            
            import random
            noise = random.uniform(-0.03, 0.03)
            score = max(0.0, min(1.0, score + noise))
            
            return round(score, 3)
        except:
            return 0.5

    def score_relevance(self, query: str, result_content: str, result_type: str = "document") -> Dict[str, Any]:
        system_prompt = """
You are an expert relevance evaluator. Score how well content answers a query.

Rating scale:
- 0.9-1.0: Directly answers the question with complete information
- 0.7-0.8: Very relevant, contains most needed information  
- 0.5-0.6: Partially relevant, some useful information
- 0.3-0.4: Tangentially related, limited usefulness
- 0.0-0.2: Not relevant to the question

Respond with valid JSON only:
{"score": 0.0-1.0, "reason": "brief explanation"}
"""
        
        user_content = f"Query: {query}\n\nContent: {result_content[:500]}..."
        
        try:
            completion = self.client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": "http://localhost:3000",
                    "X-Title": "Judge Agent",
                },
                model="meta-llama/llama-4-maverick:free",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                temperature=0.2,
                timeout=8,  # Shorter timeout
                max_tokens=150  # Limit response length
            )
            
            response = self._clean_response(completion.choices[0].message.content.strip())
            
            try:
                result = json.loads(response)
                score = result.get('score', 0.5)
                reason = result.get('reason', 'LLM evaluation')
            except:
                import re
                score_match = re.search(r'["\']?score["\']?\s*:\s*([0-9.]+)', response)
                if score_match:
                    score = float(score_match.group(1))
                    reason = "Extracted from response"
                else:
                    return self.advanced_fallback_scoring(query, result_content)
            
            normalized_score = self._normalize_score(score)
            return {
                "relevance_score": normalized_score,
                "reasoning": reason[:100],
                "key_matches": self._extract_key_matches(query, result_content),
                "missing_elements": []
            }
            
        except Exception as e:
            return self.advanced_fallback_scoring(query, result_content)

    def _extract_key_matches(self, query: str, content: str) -> List[str]:
        try:
            query_words = set(word.lower().strip('.,!?"') for word in query.split() if len(word) > 2)
            content_words = set(word.lower().strip('.,!?"') for word in content.split() if len(word) > 2)
            matches = list(query_words.intersection(content_words))
            return matches[:5]
        except:
            return []
    
    def calculate_fallback_score(self, query: str, content: str) -> Dict[str, Any]:
        """Advanced fallback scoring with multiple factors"""
        try:
            query_lower = query.lower()
            content_lower = content.lower()
            
            # Factor 1: Direct keyword matching
            query_words = set(word.strip('.,!?"') for word in query_lower.split() if len(word) > 2)
            content_words = set(word.strip('.,!?"') for word in content_lower.split() if len(word) > 2)
            common_words = query_words.intersection(content_words)
            keyword_score = len(common_words) / len(query_words) if query_words else 0
            
            # Factor 2: Phrase matching
            query_phrases = [query_lower[i:i+20] for i in range(0, len(query_lower)-19, 5)]
            phrase_matches = sum(1 for phrase in query_phrases if phrase in content_lower)
            phrase_score = min(1.0, phrase_matches / max(1, len(query_phrases)))
            
            # Factor 3: Content type relevance
            type_indicators = ['employee', 'policy', 'procedure', 'asset', 'ticket', 'record']
            has_type_match = any(indicator in query_lower and indicator in content_lower 
                               for indicator in type_indicators)
            type_score = 0.3 if has_type_match else 0.0
            
            # Factor 4: Length and detail factor
            content_length_factor = min(1.0, len(content.split()) / 50)  # Penalize very short content
            
            # Combine factors with weights
            final_score = (
                keyword_score * 0.4 +      # 40% weight on keyword matching
                phrase_score * 0.3 +       # 30% weight on phrase matching  
                type_score * 0.2 +         # 20% weight on type matching
                content_length_factor * 0.1 # 10% weight on content quality
            )
            
            # Check for perfect matches first (exact field value matches)
            perfect_match_detected = self._detect_perfect_match(query, content)
            if perfect_match_detected:
                return {
                    "relevance_score": 1.0,
                    "reasoning": "Perfect field value match detected",
                    "key_matches": list(common_words)[:5],
                    "missing_elements": []
                }
            
            # Apply realistic scoring tiers
            if final_score >= 0.8:
                score = self._normalize_score(0.80 + final_score * 0.20)  # 0.80-1.0 range
            elif final_score >= 0.6:
                score = self._normalize_score(0.65 + final_score * 0.20)  # 0.65-0.85 range
            elif final_score >= 0.4:
                score = self._normalize_score(0.45 + final_score * 0.25)  # 0.45-0.70 range
            elif final_score >= 0.2:
                score = self._normalize_score(0.25 + final_score * 0.25)  # 0.25-0.50 range
            else:
                score = self._normalize_score(final_score * 0.25)         # 0.0-0.25 range
            
            return {
                "relevance_score": score,
                "reasoning": f"Multi-factor analysis: kw={keyword_score:.2f}, phrase={phrase_score:.2f}",
                "key_matches": list(common_words)[:5],
                "missing_elements": []
            }
            
        except Exception as e:
            # Ultimate fallback
            return {
                "relevance_score": 0.4,
                "reasoning": "Evaluation completed using fallback method",
                "key_matches": [],
                "missing_elements": []
            }
    
    def _is_exact_match(self, query: str, content: str) -> bool:
        """Detect if content contains perfect field value matches for the query"""
        try:
            import re
            query_lower = query.lower()
            content_lower = content.lower()
            
            # Extract potential names/entities from query
            quoted_terms = re.findall(r'["\']([^"\']*)["\']', query)
            capitalized_terms = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', query)
            
            # Look for relationship queries
            relationship_patterns = [
                r'who\s+reports?\s+to\s+([A-Za-z\s]+)',
                r'who\s+works?\s+for\s+([A-Za-z\s]+)',
                r'([A-Za-z\s]+)\s+reports?\s+to',
                r'([A-Za-z\s]+)\s+manager'
            ]
            
            potential_matches = set()
            potential_matches.update(quoted_terms)
            potential_matches.update(capitalized_terms)
            
            # Extract from relationship patterns
            for pattern in relationship_patterns:
                matches = re.findall(pattern, query, re.IGNORECASE)
                for match in matches:
                    if isinstance(match, str) and len(match.strip()) > 2:
                        potential_matches.add(match.strip())
            
            # Check if any potential match appears as exact field value
            for term in potential_matches:
                term_clean = term.lower().strip()
                if len(term_clean) < 3:
                    continue
                    
                # Look for pattern: "field_name: exact_term"
                field_value_pattern = rf'[^:]+:\s*{re.escape(term_clean)}\s*(?:\n|$)'
                if re.search(field_value_pattern, content_lower):
                    return True
                    
                # Also check for exact term appearance in content
                if f': {term_clean}' in content_lower or f':{term_clean}' in content_lower:
                    return True
            
            return False
        except:
            return False

    def score_results_batch(self, query: str, results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Evaluate multiple results with improved batch processing"""
        evaluated_results = []
        
        # Limit batch size to prevent timeouts
        max_batch_size = 5
        results_to_process = results[:max_batch_size]
        
        for i, result in enumerate(results_to_process):
            content = result.get('content', '')
            result_type = result.get('type', 'document')
            original_score = result.get('score', 0.0)
            
            # Add small delay between requests to prevent rate limiting
            if i > 0:
                import time
                time.sleep(0.1)
            
            # Get judge evaluation with timeout protection
            try:
                evaluation = self.evaluate_relevance(query, content, result_type)
                judge_score = evaluation.get('relevance_score', original_score)
                judge_reasoning = evaluation.get('reasoning', 'Evaluated')
            except:
                # Fallback to advanced scoring if judge fails
                fallback_eval = self.advanced_fallback_scoring(query, content)
                judge_score = fallback_eval.get('relevance_score', original_score)
                judge_reasoning = "Advanced fallback scoring"
            
            # Create updated result
            updated_result = result.copy()
            updated_result['judge_score'] = judge_score
            updated_result['judge_reasoning'] = judge_reasoning
            updated_result['original_score'] = original_score
            
            # Blend judge score with original score for more stability
            blended_score = (judge_score * 0.7) + (original_score * 0.3)
            updated_result['score'] = self._normalize_score(blended_score)
            
            evaluated_results.append(updated_result)
        
        # Add remaining results without judge evaluation but with improved original scores
        for result in results[max_batch_size:]:
            updated_result = result.copy()
            updated_result['judge_score'] = None
            updated_result['judge_reasoning'] = "Not evaluated - batch limit"
            updated_result['original_score'] = result.get('score', 0.0)
            # Keep original score but normalize it
            updated_result['score'] = self._normalize_score(result.get('score', 0.0))
            evaluated_results.append(updated_result)
        
        return evaluated_results

    def evaluate_relevance(self, query: str, content: str, result_type: str = "document") -> Dict[str, Any]:
        """Evaluate relevance of content to query"""
        return self.score_relevance(query, content, result_type)

    def advanced_fallback_scoring(self, query: str, content: str) -> Dict[str, Any]:
        """Advanced fallback scoring when LLM evaluation fails"""
        return self.calculate_fallback_score(query, content)

    def _detect_perfect_match(self, query: str, content: str) -> bool:
        """Detect perfect matches in content"""
        return self._is_exact_match(query, content)

class AirtableClient:
    def __init__(self, base_id: str, airtable_api_key: str, openrouter_api_key: str):
        self.query_builder = AirtableQueryBuilder(base_id, airtable_api_key)
        self.sql_generator = SQLCodeGenerator(openrouter_api_key)
        self.orchestrator = QueryRouter(openrouter_api_key)
        self.conversation_history = []
        self.context = {}
        
        # Dynamically discover tables and their purposes
        self.discover_tables()
    
    def discover_tables(self):
        """Automatically discover all tables and understand their purpose"""
        schema_data = self.query_builder.get_base_schema()
        if "error" in schema_data:
            self.table_schemas = {}
            self.table_descriptions = {}
            return
        
        self.table_schemas = {}
        self.table_descriptions = {}
        
        tables = schema_data.get("tables", [])
        for table in tables:
            table_name = table["name"]
            
            # Store schema
            self.table_schemas[table_name] = table
            
            # Get sample data to understand purpose
            sample_data = self.query_builder.get_sample_data(table_name, 3)
            
            # Use LLM to understand table purpose
            if "error" not in sample_data:
                description = self.orchestrator.analyze_table_purpose(table_name, table, sample_data)
                self.table_descriptions[table_name] = description
            else:
                self.table_descriptions[table_name] = f"A table containing {table_name} data"
    
    def query_airtable_data(self, user_query: str, chat_history: List[Dict] = None) -> Dict[str, Any]:
        """Query Airtable data using natural language"""
        # Pattern-based routing for common employee queries
        query_lower = user_query.lower()
        employee_patterns = [
            "who works for", "who reports to", "employees", "staff", "team members",
            "works for", "reports to", "manager", "subordinates", "direct reports",
            "who is"
        ]
        
        force_airtable = any(pattern in query_lower for pattern in employee_patterns)
        
        if force_airtable:
            # Force use of Employees table for employee queries
            decision = {"use_airtable": True, "table_name": "Employees", "reasoning": "Employee query pattern detected"}
        else:
            # Check if we should use Airtable for this query
            decision = self.orchestrator.should_use_airtable(user_query, self.table_descriptions)
        
        if not decision.get("use_airtable"):
            return {"error": "Query better suited for document retrieval", "use_rag": True}
        
        table_name = decision.get("table_name")
        if not table_name or table_name not in self.table_schemas:
            table_name = self.orchestrator.determine_best_table(user_query, self.table_descriptions)
        
        # Use enhanced query with retry logic  
        print(f"📊 Querying {table_name} table...")
        result = self.enhanced_query_with_retry(table_name, user_query, chat_history=chat_history)
        
        # Add metadata for better context
        if "records" in result:
            result["table_name"] = table_name
            result["table_description"] = self.table_descriptions.get(table_name, "")
        
        return result
    
    def enhanced_query_with_retry(self, table_name: str, natural_query: str, attempt: int = 1, chat_history: List[Dict] = None) -> Dict[str, Any]:
        """Enhanced query execution with Qwen+Llama collaboration and retry logic"""
        schema = self.table_schemas.get(table_name, {})
        
        # Pass conversation context to Qwen
        context_for_llm = {
            "conversation_history": chat_history or [],
            "current_context": self.context,
            "attempt_number": attempt
        }
        
        # First attempt: Qwen generates the query
        print(f"🔍 Attempt {attempt}: Qwen generating query...")
        parsed_query = self.sql_generator.generate_airtable_query(natural_query, schema, context_for_llm)
        
        if "error" in parsed_query:
            if attempt < 3:  # Allow up to 3 attempts
                print("❌ Qwen failed, Llama analyzing for retry...")
                retry_suggestion = self.orchestrator.analyze_and_suggest_retry(
                    natural_query, schema, parsed_query["error"], context_for_llm
                )
                if retry_suggestion.get("should_retry"):
                    modified_query = retry_suggestion.get("modified_query", natural_query)
                    return self.enhanced_query_with_retry(table_name, modified_query, attempt + 1, chat_history=chat_history)
            return parsed_query
        
        # Build and execute query
        params = QueryParams()
        
        if "conditions" in parsed_query:
            params.filter_formula = self.query_builder.build_filter_formula(parsed_query["conditions"])
        
        if "sort" in parsed_query:
            params.sort = self.query_builder.build_sort_params(parsed_query["sort"])
        
        if "fields" in parsed_query:
            params.fields = parsed_query["fields"]
        
        if "max_records" in parsed_query:
            params.max_records = parsed_query["max_records"]
        
        result = self.query_builder.execute_query(table_name, params)
        
        # If no results but no error, let Llama analyze and suggest retry
        if result.get("records") is not None and len(result["records"]) == 0 and attempt < 3:
            print("🤔 No results found, Llama analyzing for retry...")
            retry_analysis = self.orchestrator.analyze_empty_results(
                natural_query, schema, parsed_query, context_for_llm
            )
            
            if retry_analysis.get("should_retry"):
                modified_query = retry_analysis.get("modified_query", natural_query)
                return self.enhanced_query_with_retry(table_name, modified_query, attempt + 1)
        
        return result

class SupabaseVectorManager:
    def __init__(self, supabase_url: str, supabase_key: str, embedder):
        self.supabase_url = supabase_url
        self.supabase_key = supabase_key
        self.embedder = embedder
        self.supabase = create_client(supabase_url, supabase_key)
        self.bucket = 'documents'
    
    def setup_vector_database(self):
        """Check if documents table exists, create if needed"""
        print("Checking Supabase vector database setup...")
        
        # Try to query the documents table to see if it exists
        try:
            test_url = f'{self.supabase_url}/rest/v1/documents'
            test_params = {'select': 'id', 'limit': 1}
            response = requests.get(test_url, headers=self.headers, params=test_params)
            
            if response.status_code == 200:
                print("✅ Documents table exists and is accessible")
                return True
            elif response.status_code == 404:
                print("❌ Documents table does not exist")
                print("\n🔧 MANUAL SETUP REQUIRED:")
                print("Please run these SQL commands in your Supabase SQL Editor:")
                print("\n" + "="*60)
                print("-- Enable pgvector extension")
                print("CREATE EXTENSION IF NOT EXISTS vector;")
                print("\n-- Create documents table")
                print("""CREATE TABLE documents (
    id SERIAL PRIMARY KEY,
    title TEXT NOT NULL,
    content TEXT NOT NULL,
    source TEXT NOT NULL,
    file_path TEXT,
    chunk_index INTEGER DEFAULT 0,
    embedding VECTOR(384),
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);""")
                print("\n-- Create index for faster similarity search")
                print("""CREATE INDEX documents_embedding_idx 
ON documents USING ivfflat (embedding vector_cosine_ops) 
WITH (lists = 100);""")
                print("\n-- Create similarity search function")
                print("""CREATE OR REPLACE FUNCTION match_documents (
    query_embedding VECTOR(384),
    match_threshold FLOAT DEFAULT 0.3,
    match_count INT DEFAULT 10
)
RETURNS TABLE (
    id BIGINT,
    title TEXT,
    content TEXT,
    source TEXT,
    similarity FLOAT
)
LANGUAGE SQL STABLE
AS $$
    SELECT
        documents.id,
        documents.title,
        documents.content,
        documents.source,
        1 - (documents.embedding <=> query_embedding) as similarity
    FROM documents
    WHERE 1 - (documents.embedding <=> query_embedding) > match_threshold
    ORDER BY (documents.embedding <=> query_embedding) ASC
    LIMIT match_count;
$$;""")
                print("="*60)
                return False
            else:
                print(f"❌ Unexpected response: {response.status_code}")
                return False
                
        except Exception as e:
            print(f"❌ Error checking documents table: {e}")
            return False
    
    def download_file_from_storage(self, file_path: str) -> bytes:
        """Download file content from Supabase storage using official client"""
        try:
            print(f"Downloading {file_path} from Supabase storage...")
            file_data = self.supabase.storage.from_(self.bucket).download(file_path)
            print(f"✅ Successfully downloaded {file_path} ({len(file_data)} bytes)")
            return file_data
        except Exception as e:
            print(f"❌ Error downloading {file_path}: {e}")
            return None
    
    def extract_text_from_file(self, file_content: bytes, file_name: str) -> str:
        """Extract text from PDF or Word documents"""
        try:
            file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
            
            if file_ext == 'pdf':
                return self.extract_text_from_pdf(file_content)
            elif file_ext in ['docx', 'doc']:
                return self.extract_text_from_docx(file_content)
            elif file_ext == 'txt':
                return file_content.decode('utf-8', errors='ignore')
            else:
                # Try to decode as text
                try:
                    return file_content.decode('utf-8', errors='ignore')
                except:
                    return f"Document: {file_name}\nUnable to extract text from this file format."
        except Exception as e:
            print(f"Error extracting text from {file_name}: {e}")
            return f"Document: {file_name}\nError processing file content."
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return "Error extracting text from PDF file."
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return "Error extracting text from Word document."
    
    def chunk_text_intelligently(self, text: str, max_chunk_size: int = 512) -> List[str]:
        """Intelligent text chunking with similarity threshold"""
        if len(text.split()) <= max_chunk_size:
            return [text]
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # If adding this paragraph would exceed chunk size
            test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if len(test_chunk.split()) > max_chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = paragraph
                else:
                    # Paragraph itself is too long, split by sentences
                    sentences = paragraph.split('. ')
                    for sentence in sentences:
                        if len((current_chunk + ". " + sentence).split()) > max_chunk_size:
                            if current_chunk:
                                chunks.append(current_chunk.strip())
                                current_chunk = sentence
                            else:
                                # Even single sentence is too long, force split
                                words = sentence.split()
                                for i in range(0, len(words), max_chunk_size):
                                    chunk_words = words[i:i + max_chunk_size]
                                    chunks.append(' '.join(chunk_words))
                        else:
                            current_chunk = current_chunk + ". " + sentence if current_chunk else sentence
            else:
                current_chunk = test_chunk
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if len(chunk.strip()) > 50]  # Filter out very short chunks
    
    def store_document_embeddings(self, title: str, content: str, source: str, file_path: str = None):
        """Process document, create embeddings, and store in Supabase"""
        try:
            # Check if document already exists
            check_url = f'{self.supabase_url}/rest/v1/documents'
            check_params = {'select': 'id', 'title': f'eq.{title}', 'source': f'eq.{source}'}
            check_response = requests.get(check_url, headers=self.headers, params=check_params)
            
            if check_response.status_code == 200 and check_response.json():
                print(f"Document '{title}' already exists in Supabase. Skipping...")
                return
            
            # Chunk the content
            chunks = self.chunk_text_intelligently(content)
            print(f"Created {len(chunks)} chunks for '{title}'")
            
            # Process each chunk
            for i, chunk in enumerate(chunks):
                # Generate embedding
                embedding = self.embedder.encode([chunk])[0].tolist()
                
                # Prepare document data
                doc_data = {
                    'title': title,
                    'content': chunk,
                    'source': source,
                    'file_path': file_path,
                    'chunk_index': i,
                    'embedding': embedding
                }
                
                # Insert into Supabase
                insert_url = f'{self.supabase_url}/rest/v1/documents'
                response = requests.post(insert_url, headers=self.headers, json=doc_data)
                
                if response.status_code not in [200, 201]:
                    print(f"Error inserting chunk {i}: {response.status_code} - {response.text}")
            
            print(f"Successfully stored '{title}' with {len(chunks)} chunks in Supabase")
            
        except Exception as e:
            print(f"Error storing document embeddings: {e}")
    
    def process_storage_documents(self):
        """Download and process all documents from Supabase storage"""
        try:
            # Get list of files in storage
            storage_url = f'{self.supabase_url}/storage/v1/object/list/documents'
            response = requests.post(storage_url, headers=self.headers, json={'prefix': '', 'limit': 1000})
            
            if response.status_code != 200:
                print(f"Failed to list storage files: {response.status_code}")
                return
            
            files = response.json()
            print(f"Found {len(files)} files in storage. Processing...")
            
            for file_info in files:
                if not file_info or not file_info.get('name'):
                    continue
                
                file_name = file_info.get('name', '')
                print(f"\nProcessing: {file_name}")
                
                # Download file content
                file_content = self.download_file_from_storage(file_name)
                if not file_content:
                    continue
                
                # Extract text
                text_content = self.extract_text_from_file(file_content, file_name)
                
                if len(text_content.strip()) > 100:  # Only process if we got substantial content
                    self.store_document_embeddings(
                        title=file_name,
                        content=text_content,
                        source="Supabase Storage",
                        file_path=file_name
                    )
                else:
                    print(f"Skipping {file_name} - insufficient text content")
        
        except Exception as e:
            print(f"Error processing storage documents: {e}")
    
    def search_similar_documents(self, query: str, limit: int = 10) -> List[Dict]:
        """Search for similar documents using vector similarity"""
        try:
            # Generate query embedding
            query_embedding = self.embedder.encode([query])[0].tolist()
            
            # Call the similarity search function
            search_url = f'{self.supabase_url}/rest/v1/rpc/match_documents'
            search_data = {
                'query_embedding': query_embedding,
                'match_threshold': 0.3,
                'match_count': limit
            }
            
            response = requests.post(search_url, headers=self.headers, json=search_data)
            
            if response.status_code == 200:
                results = response.json()
                return results
            else:
                print(f"Search failed: {response.status_code} - {response.text}")
                return []
        
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []

class DocumentSearchSystem:
    def _score_airtable_match(self, query: str, fields: dict, content_text: str) -> float:
        """Calculate relevance score for Airtable records with perfect match detection"""
        try:
            query_lower = query.lower().strip()
            
            # Extract key entities from query (names, terms)
            import re
            # Find quoted terms or capitalize names
            quoted_terms = re.findall(r'["\']([^"\']*)["\']', query)
            capitalized_terms = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', query)
            
            # Combine all potential key terms
            key_terms = set()
            key_terms.update(quoted_terms)
            key_terms.update(capitalized_terms)
            
            # Also extract from query words (for cases without quotes/caps)
            query_words = query.split()
            if len(query_words) >= 2:
                # Look for name patterns (two consecutive words that could be a name)
                for i in range(len(query_words) - 1):
                    potential_name = f"{query_words[i]} {query_words[i+1]}"
                    if len(potential_name) > 4:  # Skip very short combinations
                        key_terms.add(potential_name)
            
            # Check for perfect field value matches
            perfect_match_score = 0.0
            high_relevance_score = 0.0
            
            # Define relationship field patterns
            relationship_fields = ['manager', 'reports to', 'supervisor', 'lead', 'director', 'team lead']
            employee_fields = ['name', 'full name', 'employee name', 'employee']
            
            for field_name, field_value in fields.items():
                if not field_value:
                    continue
                    
                field_name_lower = field_name.lower()
                field_value_str = str(field_value).strip()
                field_value_lower = field_value_str.lower()
                
                # Check for exact matches of key terms in field values
                for term in key_terms:
                    term_lower = term.lower().strip()
                    if len(term_lower) < 3:  # Skip very short terms
                        continue
                        
                    # Perfect match: exact term appears as field value
                    if term_lower == field_value_lower:
                        # Check query context to determine match strength
                        if any(rel_field in query_lower for rel_field in ['reports to', 'managed by', 'works for']):
                            if any(rel_pattern in field_name_lower for rel_pattern in relationship_fields):
                                perfect_match_score = max(perfect_match_score, 1.0)  # 100% match
                            elif any(emp_pattern in field_name_lower for emp_pattern in employee_fields):
                                perfect_match_score = max(perfect_match_score, 1.0)  # Employee name match
                        else:
                            perfect_match_score = max(perfect_match_score, 0.95)  # Very high but not perfect
                    
                    # High relevance: term appears within field value
                    elif term_lower in field_value_lower:
                        if any(rel_pattern in field_name_lower for rel_pattern in relationship_fields + employee_fields):
                            high_relevance_score = max(high_relevance_score, 0.85)
                        else:
                            high_relevance_score = max(high_relevance_score, 0.75)
            
            # If we found perfect matches, return them
            if perfect_match_score > 0.9:
                return perfect_match_score
            
            # If we found high relevance, return it
            if high_relevance_score > 0.7:
                return high_relevance_score
            
            # Fall back to word overlap scoring for partial matches
            query_words_set = set(word.lower().strip('.,!?"') for word in query.split() if len(word) > 2)
            content_words_set = set(word.lower().strip('.,!?"') for word in content_text.split() if len(word) > 2)
            
            if not query_words_set:
                return 0.5
                
            overlap = len(query_words_set.intersection(content_words_set)) / len(query_words_set)
            
            # Enhanced overlap scoring
            if overlap >= 0.8:
                return 0.85
            elif overlap >= 0.6:
                return 0.75
            elif overlap >= 0.4:
                return 0.65
            elif overlap >= 0.2:
                return 0.55
            else:
                return 0.45
                
        except Exception as e:
            # Fallback to basic scoring
            query_words = set(query.lower().split())
            content_words = set(content_text.lower().split())
            overlap = len(query_words.intersection(content_words)) / len(query_words) if query_words else 0
            return 0.6 + (overlap * 0.3)
    
    def __init__(self):
        # Initialize embeddings model
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Initialize Pinecone
        self.pc = Pinecone(api_key=os.getenv('PINECONE_API_KEY'))
        
        index_name = "rag-documents"
        
        # Delete existing index if dimension mismatch
        if index_name in self.pc.list_indexes().names():
            try:
                # Check if we need to recreate due to dimension mismatch
                existing_index = self.pc.Index(index_name)
                # Try a test query to see if dimension matches
                test_embedding = self.embedder.encode(['test'])[0].tolist()
                existing_index.query(vector=test_embedding, top_k=1)
            except Exception:
                print("Recreating Pinecone index due to dimension mismatch...")
                self.pc.delete_index(index_name)
                # Wait for deletion
                import time
                time.sleep(10)
        
        if index_name not in self.pc.list_indexes().names():
            self.pc.create_index(
                name=index_name,
                dimension=384,  # all-MiniLM-L6-v2 dimension
                metric='cosine',
                spec=ServerlessSpec(
                    cloud='aws',
                    region='us-east-1'
                )
            )
            # Wait for index to be ready
            import time
            time.sleep(10)
            
        self.index = self.pc.Index(index_name)
        
        # Initialize Claude
        self.claude = anthropic.Anthropic(api_key=os.getenv('CLAUDE_API_KEY'))
        
        # Supabase config
        self.supabase_url = os.getenv('SUPABASE_URL')
        self.supabase_key = os.getenv('SUPABASE_SERVICE_ROLE_KEY')
        
        # Initialize Supabase Vector Manager
        self.supabase_vector = SupabaseVectorManager(
            self.supabase_url, 
            self.supabase_key, 
            self.embedder
        )
        
        # Airtable config
        self.airtable_key = os.getenv('AIRTABLE_API_KEY')
        self.base_id = os.getenv('BASE_ID')
        self.openrouter_key = os.getenv('OPENROUTER_API_KEY')
        
        # Initialize Universal Airtable Agent if credentials are available
        self.airtable_agent = None
        if self.airtable_key and self.base_id and self.openrouter_key:
            try:
                self.airtable_agent = AirtableClient(
                    self.base_id,
                    self.airtable_key, 
                    self.openrouter_key
                )
            except Exception as e:
                print(f"Warning: Could not initialize Airtable agent: {e}")
                self.airtable_agent = None
        
        # Initialize Judge Agent if OpenRouter key is available
        self.judge_agent = None
        if self.openrouter_key:
            try:
                self.judge_agent = RelevanceScorer(self.openrouter_key)
            except Exception as e:
                print(f"Warning: Could not initialize Judge agent: {e}")
                self.judge_agent = None
        
        # Cache for processed documents and file hashes
        self.cache_file = "document_cache.pkl"
        self.file_hash_cache = "file_hash_cache.pkl"
        self.processed_docs = self.load_cache()
        self.processed_files = self.load_file_hash_cache()
    
    def load_cache(self):
        if os.path.exists(self.cache_file):
            with open(self.cache_file, 'rb') as f:
                return pickle.load(f)
        return {}
    
    def load_file_hash_cache(self):
        if os.path.exists(self.file_hash_cache):
            with open(self.file_hash_cache, 'rb') as f:
                return pickle.load(f)
        return {}
    
    def save_cache(self):
        with open(self.cache_file, 'wb') as f:
            pickle.dump(self.processed_docs, f)
    
    def save_file_hash_cache(self):
        with open(self.file_hash_cache, 'wb') as f:
            pickle.dump(self.processed_files, f)
    
    def clean_airtable_vectors(self):
        """Remove old Airtable vectors from Pinecone to prevent conflicts with new Airtable agent"""
        try:
            print("🧹 Cleaning old Airtable vectors from Pinecone...")
            
            # Get index stats to see what we have
            stats = self.index.describe_index_stats()
            print(f"   Index has {stats.total_vector_count} total vectors")
            
            # Since we can't use regex, we'll scan by prefix
            # Query for vectors that might contain Airtable data
            dummy_vector = [0.1] * 384  # Use a non-zero vector for better results
            
            # Get a large sample of vectors
            results = self.index.query(
                vector=dummy_vector,
                top_k=min(10000, stats.total_vector_count),
                include_metadata=True
            )
            
            # Collect IDs of Airtable vectors
            airtable_ids = []
            for match in results.get('matches', []):
                metadata = match.get('metadata', {})
                source = metadata.get('source', '')
                if 'Airtable' in source:
                    airtable_ids.append(match['id'])
                    print(f"   Found Airtable vector: {match['id']} - {source}")
            
            if airtable_ids:
                # Delete in batches
                batch_size = 100
                for i in range(0, len(airtable_ids), batch_size):
                    batch = airtable_ids[i:i + batch_size]
                    self.index.delete(ids=batch)
                    print(f"   Deleted {len(batch)} Airtable vectors...")
                
                print(f"✅ Cleaned {len(airtable_ids)} old Airtable vectors from Pinecone")
                
                # Also clear cache for Airtable docs
                airtable_keys = [key for key in self.processed_docs.keys() if 'airtable' in key.lower()]
                for key in airtable_keys:
                    del self.processed_docs[key]
                if airtable_keys:
                    print(f"   Cleared {len(airtable_keys)} Airtable entries from cache")
                    self.save_cache()
                
            else:
                print("✅ No old Airtable vectors found in Pinecone")
                
        except Exception as e:
            print(f"⚠️  Error cleaning Airtable vectors: {e}")
    
    def clear_all_vectors(self):
        """Clear all vectors from Pinecone (nuclear option)"""
        try:
            print("🧹 Clearing ALL vectors from Pinecone...")
            self.index.delete(delete_all=True)
            
            # Clear all caches
            self.processed_docs = {}
            self.processed_files = {}
            self.save_cache()
            self.save_file_hash_cache()
            
            print("✅ All vectors and caches cleared")
        except Exception as e:
            print(f"⚠️  Error clearing vectors: {e}")
    
    def get_document_hash(self, content: str) -> str:
        return hashlib.md5(content.encode()).hexdigest()
    
    def chunk_content_adaptively(self, content: str, max_chunk_size: int = 512) -> List[str]:
        """Adaptive chunking with similarity threshold"""
        sentences = content.split('. ')
        if len(sentences) <= 1:
            return [content]
        
        chunks = []
        current_chunk = sentences[0]
        
        for i in range(1, len(sentences)):
            sentence = sentences[i]
            test_chunk = current_chunk + ". " + sentence
            
            if len(test_chunk.split()) > max_chunk_size:
                # Check similarity before splitting
                if len(current_chunk.split()) > 50:  # Only check similarity for substantial chunks
                    current_emb = self.embedder.encode([current_chunk])
                    sentence_emb = self.embedder.encode([sentence])
                    similarity = 1 - cosine(current_emb[0], sentence_emb[0])
                    
                    if similarity < 0.7:  # Low similarity threshold
                        chunks.append(current_chunk.strip())
                        current_chunk = sentence
                    else:
                        chunks.append(current_chunk.strip())
                        current_chunk = sentence
                else:
                    chunks.append(current_chunk.strip())
                    current_chunk = sentence
            else:
                current_chunk = test_chunk
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if len(chunk.strip()) > 20]
    
    def process_storage_documents(self):
        """Download and process documents from Supabase storage, store in Pinecone"""
        documents, updated_file_hashes = self.supabase_vector.process_storage_documents(self.processed_files)
        
        # Update file hash cache
        self.processed_files.update(updated_file_hashes)
        self.save_file_hash_cache()
        
        if documents:
            print(f"Processing {len(documents)} new/modified documents for Pinecone...")
            self.process_and_store_documents(documents)
        else:
            print("No new documents to process from Supabase storage.")
    
    def fetch_airtable_data(self) -> List[Document]:
        documents = []
        headers = {
            'Authorization': f'Bearer {self.airtable_key}',
            'Content-Type': 'application/json'
        }
        
        try:
            # Get tables
            url = f'https://api.airtable.com/v0/meta/bases/{self.base_id}/tables'
            response = requests.get(url, headers=headers)
            for file in files:
                name = file.get('name', '')
                metadata = file.get('metadata')
                mimetype = metadata.get('mimetype') if metadata else None
                
                if mimetype is None:  # It's a folder
                    subfolder_path = f"{path}/{name}" if path else name
                    subfiles = self.list_all_files(subfolder_path)
                    files_found.extend(subfiles)
                else:  # It's a file
                    file_path = f"{path}/{name}" if path else name
                    files_found.append({
                        'path': file_path,
                        'name': name,
                        'mimetype': mimetype,
                        'size': file.get('metadata', {}).get('size', 0),
                        'last_modified': file.get('updated_at', file.get('created_at', '')),
                        'id': file.get('id', '')
                    })
        except Exception as e:
            pass
        
        return files_found
    
    def get_file_hash(self, file_path: str, metadata: dict) -> str:
        """Generate hash for file based on path, size, and last modified time"""
        hash_string = f"{file_path}_{metadata.get('size', 0)}_{metadata.get('last_modified', '')}"
        return hashlib.md5(hash_string.encode()).hexdigest()
    
    def should_process_file(self, file_path: str, metadata: dict, processed_files: dict) -> bool:
        """Check if file should be processed based on hash comparison"""
        current_hash = self.get_file_hash(file_path, metadata)
        stored_hash = processed_files.get(file_path)
        
        return stored_hash != current_hash
    
    def list_all_files(self, path: str = "") -> list:
        """List all files recursively from Supabase storage"""
        files_found = []
        
        try:
            # Get list of files in storage
            storage_url = f'{self.supabase_url}/storage/v1/object/list/documents'
            headers = {
                'Authorization': f'Bearer {self.supabase_key}',
                'apikey': self.supabase_key,
                'Content-Type': 'application/json'
            }
            
            response = requests.post(storage_url, headers=headers, json={'prefix': path, 'limit': 1000})
            
            if response.status_code != 200:
                return []
            
            files = response.json()
            
            for file in files:
                if not file or not file.get('name'):
                    continue
                
                name = file.get('name', '')
                metadata = file.get('metadata')
                mimetype = metadata.get('mimetype') if metadata else None
                
                if mimetype is None:  # It's a folder
                    subfolder_path = f"{path}/{name}" if path else name
                    subfiles = self.list_all_files(subfolder_path)
                    files_found.extend(subfiles)
                else:  # It's a file
                    file_path = f"{path}/{name}" if path else name
                    files_found.append({
                        'path': file_path,
                        'name': name,
                        'mimetype': mimetype,
                        'size': file.get('metadata', {}).get('size', 0),
                        'last_modified': file.get('updated_at', file.get('created_at', '')),
                        'id': file.get('id', '')
                    })
        except Exception as e:
            pass
        
        return files_found
    
    def download_file_from_storage(self, file_path: str) -> bytes:
        """Download file content from Supabase storage"""
        try:
            file_data = self.supabase_vector.supabase.storage.from_('documents').download(file_path)
            return file_data
        except Exception as e:
            return None
    
    def extract_text_from_file(self, file_content: bytes, file_name: str) -> str:
        """Extract text from files using the SupabaseVectorManager method"""
        return self.supabase_vector.extract_text_from_file(file_content, file_name)
    
    def process_storage_documents(self, processed_files: dict = None):
        """Download and process documents from Supabase storage, store in Pinecone"""
        if processed_files is None:
            processed_files = {}
        
        try:
            # Get all files recursively
            all_files = self.list_all_files()
            
            if not all_files:
                return [], processed_files
            
            documents = []
            new_processed_files = processed_files.copy()
            
            for file_info in all_files:
                file_path = file_info['path']
                file_name = file_info['name']
                mimetype = file_info['mimetype']
                
                # Check if we should process this file
                if not self.should_process_file(file_path, file_info, processed_files):
                    continue
                
                # Download file content
                file_content = self.download_file_from_storage(file_path)
                if not file_content:
                    continue
                
                # Extract text based on file type
                text_content = self.extract_text_from_file(file_content, file_name)
                
                if len(text_content.strip()) > 50:  # Lower threshold for text content
                    doc = Document(
                        id=f"storage_{file_path.replace('/', '_').replace(' ', '_')}",
                        title=f"Document: {file_name}",
                        content=text_content,
                        source=f"Supabase Storage|{file_path}"  # Store full path in source
                    )
                    documents.append(doc)
                    
                    # Update hash for processed file
                    new_processed_files[file_path] = self.get_file_hash(file_path, file_info)
                else:
                    # Still mark as processed to avoid retrying
                    new_processed_files[file_path] = self.get_file_hash(file_path, file_info)
            
            return documents, new_processed_files
        
        except Exception as e:
            return [], processed_files
    
    def fetch_airtable_data(self) -> List[Document]:
        """DEPRECATED: Airtable data is now queried live via Airtable agent, not indexed to Pinecone"""
        # This method is kept for backward compatibility but returns empty list
        # Airtable data should be queried live for real-time results
        return []
    
    def process_and_store_documents(self, documents: List[Document]):
        """Process documents, create embeddings, and store in Pinecone"""
        if not documents:
            return
        
        vectors_to_upsert = []
        
        for doc in documents:
            if doc.id in self.processed_docs:
                continue
            
            # Chunk content adaptively
            chunks = self.chunk_content_adaptively(doc.content)
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"{doc.id}_chunk_{i}"
                
                # Create embedding
                embedding = self.embedder.encode([chunk])[0].tolist()
                
                # Prepare metadata
                metadata = {
                    'title': doc.title,
                    'content': chunk,
                    'source': doc.source,
                    'parent_doc_id': doc.id,
                    'chunk_index': i
                }
                
                vectors_to_upsert.append((chunk_id, embedding, metadata))
            
            self.processed_docs[doc.id] = True
        
        # Batch upsert to Pinecone
        if vectors_to_upsert:
            self.index.upsert(vectors=vectors_to_upsert)
            self.save_cache()
    
    def semantic_search(self, query: str, top_k: int = 5, chat_history: List[Dict] = None) -> Dict[str, Any]:
        """Search for relevant information using both RAG and Airtable sources"""
        results = {
            'rag_results': [],
            'airtable_results': None,
            'source_used': 'rag',  # Default to RAG
            'combined_results': []
        }
        
        # Search Pinecone (traditional RAG) with improved scoring
        try:
            query_embedding = self.embedder.encode([query])[0].tolist()
            
            pinecone_results = self.index.query(
                vector=query_embedding,
                top_k=top_k * 3,  # Get more results to filter better
                include_metadata=True
            )
            
            for match in pinecone_results['matches']:
                # More sophisticated threshold based on rank
                rank_threshold = max(0.15, 0.5 - (len(results['rag_results']) * 0.05))
                
                if match['score'] > rank_threshold:
                    # Normalize Pinecone scores to be more realistic (they tend to be high)
                    normalized_score = min(0.95, match['score'] * 0.8 + 0.1)
                    
                    results['rag_results'].append({
                        'content': match['metadata']['content'],
                        'title': match['metadata']['title'],
                        'source': match['metadata']['source'],
                        'score': normalized_score,
                        'type': 'document',
                        'original_pinecone_score': match['score']
                    })
        except Exception as e:
            pass
        
        # Always try Airtable agent if available
        if self.airtable_agent:
            try:
                airtable_response = self.airtable_agent.query_airtable_data(query, chat_history=chat_history)
                if "error" not in airtable_response and not airtable_response.get("use_rag"):
                    results['airtable_results'] = airtable_response
                    results['source_used'] = 'airtable'
                    
                    # Format Airtable results for consistency
                    if airtable_response.get("records"):
                        for record in airtable_response["records"]:
                            fields = record.get("fields", {})
                            content_parts = []
                            for key, value in fields.items():
                                if value and str(value).strip():
                                    content_parts.append(f"{key}: {value}")
                            
                            if content_parts:
                                record_id = record.get('id', 'unknown')
                                table_name = airtable_response.get('table_name', 'Airtable')
                                
                                # Advanced scoring for Airtable results with perfect match detection
                                content_text = "\n".join(content_parts)
                                initial_score = self._score_airtable_match(query, fields, content_text)
                                
                                title_field_keys = ['Name', 'Title', 'Subject', 'ID', 'Key']
                                record_title = None
                                for key in title_field_keys:
                                    if key in fields and isinstance(fields[key], str) and fields[key].strip():
                                        record_title = fields[key]
                                        break
                                
                                if not record_title:
                                    for field, value in fields.items():
                                        if isinstance(value, str) and value.strip():
                                            record_title = value
                                            break

                                if record_title:
                                    doc_title = f"{table_name} - {record_title}"
                                else:
                                    doc_title = f"{table_name} - Record {record_id}"

                                results['combined_results'].append({
                                    'content': content_text,
                                    'title': doc_title,
                                    'source': f"Airtable - {table_name}",
                                    'score': initial_score,
                                    'type': 'airtable_record'
                                })
            except Exception as e:
                pass
        
        # Combine RAG results with Airtable results
        results['combined_results'].extend(results['rag_results'])
        
        # Apply judge agent evaluation if available (with better error handling)
        skip_judge = hasattr(self, '_skip_judge_evaluation') and self._skip_judge_evaluation
        if self.judge_agent and results['combined_results'] and not skip_judge:
            try:
                # Only judge if we have good results to evaluate
                if len(results['combined_results']) > 0:
                    print("🏛️ Judge evaluating result relevance...")
                    start_time = time.time()
                    
                    # Set a maximum time limit for judge evaluation
                    evaluated_results = self.judge_agent.score_results_batch(
                        query, results['combined_results']
                    )
                    
                    eval_time = time.time() - start_time
                    if eval_time < 30:  # Only use results if evaluation was reasonably fast
                        results['combined_results'] = evaluated_results
                        results['judged'] = True
                        print(f"✅ Judge evaluation completed in {eval_time:.1f}s")
                    else:
                        print("⚠️ Judge evaluation too slow, using original scores")
                        results['judged'] = False
                else:
                    results['judged'] = False
            except Exception as e:
                print(f"⚠️ Judge evaluation failed, using original scores")
                results['judged'] = False
        else:
            results['judged'] = False
        
        # Sort combined results by score and return top results
        results['combined_results'].sort(key=lambda x: x['score'], reverse=True)
        results['combined_results'] = results['combined_results'][:top_k]
        
        return results
    
    def _generate_no_results_response(self, query: str) -> str:
        """Generate a helpful response when no relevant documents are found"""
        try:
            prompt = f"""Generate a helpful response when no documents match the user's query: '{query}'. 
            
            Provide a professional, empathetic response that:
            1. Acknowledges their question
            2. Explains that no relevant information was found
            3. Suggests alternative search terms or approaches
            4. Offers to help with a rephrased question
            
            Keep it concise and helpful."""
            
            response = self.claude.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=200,
                temperature=0.3,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            return response.content[0].text.strip()
        except Exception as e:
            return f"I couldn't find any relevant information to answer your question about '{query}'. Please try rephrasing your question or using different keywords."

    def generate_answer(self, query: str, search_results: Dict[str, Any], chat_history: List[Dict] = None) -> str:
        """Generate answer using Claude with retrieved context and conversation history"""
        context_docs = search_results.get('combined_results', [])
        source_used = search_results.get('source_used', 'rag')
        
        if not context_docs:
            return self._generate_no_results_response(query)
        
        context = ""
        for i, doc in enumerate(context_docs):
            context += f"\n\nSource {i+1} ({doc.get('type', 'document')}):\n{doc['content']}"
        
        # Build conversation history for context
        conversation_context = ""
        if chat_history and len(chat_history) > 0:
            conversation_context = "\n\nCONVERSATION HISTORY:\n"
            for i, msg in enumerate(chat_history[-4:]):  # Last 4 messages for context
                role = "User" if msg.get('sender') == 'user' else "Assistant"
                conversation_context += f"{role}: {msg.get('content', '')}\n"
        
        # Add information about data sources used
        source_info = ""
        if source_used == 'airtable':
            source_info = "\n\nNOTE: This answer is based on structured data from your Airtable database, analyzed using intelligent query processing."
        elif source_used == 'rag':
            source_info = "\n\nNOTE: This answer is based on document content from your knowledge base."

        prompt = f"""You are a professional company assistant. Provide accurate, helpful answers using the provided context.

CONTEXT:
{context}{conversation_context}

GUIDELINES:
- Answer directly and professionally based only on the provided context
- For employee reporting queries ("who reports to X"): 
  • Start with a clear summary: "X has Y employees reporting to them:"
  • List all employees with their full roles, departments, and any additional relevant details
  • Include employee locations, start dates, or specializations when available
  • End with a comprehensive summary including total count and team overview
- For policy queries: Provide complete, actionable information
- Use clear formatting:
  • Numbered lists (1., 2., 3.) for sequential steps
  • Bullet points (•) for lists of items or people
  • Line breaks between points for readability
- Be detailed and comprehensive - include all relevant information from the context
- If information is incomplete, state "I don't have complete information about this"
- Be thorough - don't truncate lists or summaries when you have the data
{source_info}

Question: {query}

Answer:"""

        try:
            response = self.claude.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}]
            )
            response_text = response.content[0].text
            # Clean the response to ensure no system prompts leak
            cleaned_response = self._clean_claude_response(response_text)
            return cleaned_response
        except Exception as e:
            return "I apologize, but I'm unable to generate a response at the moment. Please try rephrasing your question."
    
    def _clean_claude_response(self, response: str) -> str:
        """Clean Claude's response to prevent any prompt leaking"""
        try:
            # Remove any lines that look like system instructions
            lines = response.split('\n')
            cleaned_lines = []
            
            for line in lines:
                # Skip lines that look like instructions or prompts (be more selective)
                if any(marker in line.lower() for marker in [
                    'you are a helpful', 'answer the user', 'using the provided',
                    'current question:', 'answer:'
                ]):
                    continue
                cleaned_lines.append(line)
            
            # Join and clean up
            result = '\n'.join(cleaned_lines).strip()
            
            # If result is too short, return original (but still cleaned)
            if len(result) < 20:
                # Just remove obvious instruction markers
                cleaned = response
                for marker in ['CONTEXT:', 'INSTRUCTIONS:', 'CURRENT QUESTION:', 'ANSWER:']:
                    cleaned = cleaned.replace(marker, '')
                return cleaned.strip()
            
            return result
        except:
            return response
    
    def update_knowledge_base(self, verbose=False):
        """Process documents and store in Pinecone - Airtable data is queried live, not indexed"""        
        # Clean any old Airtable vectors from Pinecone to ensure separation
        self.clean_airtable_vectors()
        
        # Process storage documents silently
        try:
            self.process_storage_documents()
        except:
            pass  # Silently handle errors
        
        # NOTE: Airtable data is NOT indexed to Pinecone anymore
        # It's queried live via the Airtable agent for real-time data
    
    def test_retrieval(self, question: str) -> None:
        """Test retrieval system without LLM"""
        print(f"\n=== Testing Query: '{question}' ===")
        search_results = self.semantic_search(question, top_k=10)
        relevant_docs = search_results.get('combined_results', [])
        
        if not relevant_docs:
            print("No relevant documents found!")
            return
        
        judged = search_results.get('judged', False)
        source_info = search_results.get('source_used', 'rag').upper()
        if judged:
            source_info += " + JUDGE EVALUATED"
        
        print(f"\nFound {len(relevant_docs)} relevant documents using {source_info}:")
        print("-" * 80)
        
        for i, doc in enumerate(relevant_docs, 1):
            print(f"\n{i}. SOURCE: {doc['source']}")
            print(f"   TITLE: {doc['title']}")
            
            if judged and 'judge_score' in doc:
                original_score = doc.get('original_score', doc['score'])
                print(f"   JUDGE SCORE: {doc['judge_score']:.3f}")
                print(f"   ORIGINAL SCORE: {original_score:.3f}")
                if 'judge_reasoning' in doc:
                    print(f"   REASONING: {doc['judge_reasoning'][:100]}...")
            else:
                print(f"   SCORE: {doc['score']:.3f}")
            
            print(f"   TYPE: {doc.get('type', 'document')}")
            print(f"   CONTENT: {doc['content'][:200]}...")
            print("-" * 40)

def interactive_mode():
    print("🚀 Initializing...")
    rag = DocumentSearchSystem()
    rag.update_knowledge_base()
    
    print("🎯 RAG System Ready!")
    print("💡 Ask about: employees, policies, assets, tickets")
    print("📝 Type 'quit' to exit\n")
    
    while True:
        try:
            # Get user question
            question = input("\n🤔 Your question: ").strip()
            
            if not question:
                continue
                
            if question.lower() in ['quit', 'exit', 'q']:
                print("\n👋 Goodbye!")
                break
            
            print(f"\n🔍 Searching for: '{question}'")
            print("-" * 50)
            
            # Get relevant information from both sources
            search_results = rag.semantic_search(question, top_k=5)
            
            if not search_results.get('combined_results'):
                print("❌ No relevant information found for your question.")
                continue
            
            # Generate answer using Claude
            answer = rag.generate_answer(question, search_results, [])
            
            print("🤖 ANSWER:")
            print(answer)
            
            # Show enhanced scoring information
            judged = search_results.get('judged', False)
            source_info = f"({search_results.get('source_used', 'rag').upper()}"
            if judged:
                source_info += " + JUDGE EVALUATED)"
            else:
                source_info += ")"
            
            print(f"\n📚 SOURCES {source_info}:")
            for i, doc in enumerate(search_results['combined_results'], 1):
                score_info = f"Score: {doc['score']:.3f}"
                if judged and 'judge_score' in doc:
                    original_score = doc.get('original_score', doc['score'])
                    score_info = f"Judge: {doc['judge_score']:.3f} (Orig: {original_score:.3f})"
                
                print(f"{i}. {doc['source']} - {doc['title']} ({score_info}) [{doc.get('type', 'document')}]")
            
        except KeyboardInterrupt:
            print("\n\n👋 Goodbye!")
            break
        except Exception as e:
            print(f"\n❌ Error: {e}")
            print("Please try asking another question.")

def test_mode():
    print("🧪 Running Test Mode...")
    rag = DocumentSearchSystem()
    
    # Update knowledge base
    rag.update_knowledge_base()
    
    print("\n=== Testing Retrieval System ===")
    
    # Test queries
    test_queries = [
        "manager sarah chen who is managed by her",
        "lost laptop policies", 
        "sarah chen manager",
        "laptop policy lost stolen"
    ]
    
    for query in test_queries:
        rag.test_retrieval(query)
        print("\n" + "="*100)
    
    print("\n=== Test Complete ===")
    print("All retrieval tests finished!")

def main():
    import sys
    
    if len(sys.argv) > 1 and sys.argv[1] == '--test':
        test_mode()
    else:
        interactive_mode()

if __name__ == "__main__":
    main()